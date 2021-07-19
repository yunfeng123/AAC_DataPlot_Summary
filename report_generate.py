import pandas as pd
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

# sections对应文档中的“节”
sec = document.sections[0]

# 页眉设置
sec.header_distance = Cm(0.4)   #页眉距离顶端距离
paragraph = sec.header.paragraphs[0]
run_header = paragraph.add_run('TestDataPlotSummary developed by AAC Test Team')
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_header.font.name = u'Calibri'
run_header.font.size = Pt(8)
run_header._element.rPr.rFonts.set(qn('w:eastAsia'), u'Calibri')
run_header.font.color.rgb = RGBColor(0,0,0)
run_header.font.italic = True


# 以下依次设置左、右、上、下页面边距
distance1 = Inches(0.3)
distance2 = Inches(0)
sec.left_margin = distance1
sec.right_margin = distance1
sec.top_margin = distance2
sec.bottom_margin = distance2

# 设置页面的宽度和高度(A4)
sec.page_width = Inches(11.69)
sec.page_height = Inches(8.27)

# 第一页主标题设置
page_title = document.add_heading("", level=1)
page_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_pt = page_title.add_run('\n\n\n' + 'TestData Plot Summary Report')
run_pt.font.name = u'Helvetica'
run_pt.font.size = Pt(40.5)
run_pt._element.rPr.rFonts.set(qn('w:eastAsia'), u'Helvetica')
run_pt.font.color.rgb = RGBColor(0, 0, 0)
run_pt.font.bold = True


pic_path = 'D:\\Python\\pythonProject\\venv\\pic'
list_pic = os.listdir(pic_path)

# pic_name 为图片名称拆分组成的DataFrame（分组名,Config,图类），columns为图片名称
pic_name = pd.DataFrame()
for i in list_pic:
    if os.path.isfile(pic_path + '\\' + i):
        list_temp = i.split('-')
        series_temp = pd.Series(list_temp[0:3])
        pic_name = pd.concat([pic_name,series_temp], axis=1)
    else:
        list_pic.remove(i)
pic_name.columns = list_pic

pic_name.sort_values(by=[2,0,1], axis=1, inplace=True)  # 排序，先按照图类别，然后是组的类别，最后是Config
pic_name.to_csv('pic_name.csv')

flag_index = pd.Series(['','',''])
for i in pic_name.columns:
    if pic_name[i][0] != flag_index[0]:
        document.add_page_break()
        page_title = document.add_heading('', level=1)
        page_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_pt = page_title.add_run(pic_name[i][0])
        run_pt.font.name = u'Helvetica'
        run_pt.font.size = Pt(18)
        run_pt._element.rPr.rFonts.set(qn('w:eastAsia'), u'Helvetica')
        run_pt.font.color.rgb = RGBColor(0, 0, 0)
        run_pt.italic = True
        run_pt.font.bold = True

        pr_pic = document.add_paragraph()
        run_pict = pr_pic.add_run()

    if pic_name[i][1] != flag_index[1]:
        page_title_2nd = document.add_heading('', level=2)
        page_title_2nd.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_ct = page_title_2nd.add_run(pic_name[i][1])
        run_ct.font.name = u'Helvetica'
        run_ct.font.size = Pt(12)
        run_ct._element.rPr.rFonts.set(qn('w:eastAsia'), u'Helvetica')
        run_ct.font.color.rgb = RGBColor(0, 0, 0)

        pr_pic = document.add_paragraph()
        run_pict = pr_pic.add_run()

    run_pict.add_picture(pic_path + '\\' + i,height=Inches(2))
    flag_index = pic_name[i]

document.save('a.docx')