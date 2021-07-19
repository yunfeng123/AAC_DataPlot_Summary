import pandas as pd
import xlwings as xw
import csv_process
import Histogram
import Boxplot
import Sweep_1 as Sweep
from tkinter import *
from tkinter import filedialog
import tkinter.messagebox
import os
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def run_csv():
    default_dir = r"文件路径"
    filepath_csv = filedialog.askopenfilename(title=u'选择CSV文件', initialdir=(os.path.expanduser(default_dir)))
    text_CSV.delete(0, END)
    text_CSV.insert(END, filepath_csv)


def run_ini():
    default_dir = r"文件路径"
    global filepath_ini
    filepath_ini = filedialog.askopenfilename(title=u'选择配置文件', initialdir=(os.path.expanduser(default_dir)))
    text_ini.delete(0, END)
    text_ini.insert(END, filepath_ini)


def run_pic():
    default_dir = r"文件路径"
    filepath_pic = filedialog.askdirectory(title=u'选择图片路径', initialdir=(os.path.expanduser(default_dir)))
    v.set(filepath_pic)


def run_report():
    pic_path = v.get()
    list_pic = os.listdir(pic_path)

    # pic_name 为图片名称拆分组成的DataFrame（分组名,Config,图类），columns为图片名称
    pic_name = pd.DataFrame()
    for i in list_pic:
        if os.path.isfile(pic_path + '\\' + i):
            list_temp = i.split('-')
            list_temp[2] = int(list_temp[2])
            series_temp = pd.Series(list_temp[0:4])
            pic_name = pd.concat([pic_name, series_temp], axis=1)
        else:
            list_pic.remove(i)
    pic_name.columns = list_pic

    pic_name.sort_values(by=[3, 2, 0, 1], axis=1, inplace=True)  # 排序，先按照station，再按照图类别，然后是组的类别，最后是Config
    # pic_name.to_csv('pic_name.csv')

    # sections对应文档中的“节”
    document = Document()
    sec = document.sections[0]

    # 页眉设置
    sec.header_distance = Cm(0.3)  # 页眉距离顶端距离
    paragraph = sec.header.paragraphs[0]
    run_header = paragraph.add_run('TestDataPlotSummary developed by AAC Test Team')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_header.font.name = u'Calibri'
    run_header.font.size = Pt(8)
    run_header._element.rPr.rFonts.set(qn('w:eastAsia'), u'Calibri')
    run_header.font.color.rgb = RGBColor(0, 0, 0)
    run_header.font.italic = True

    # 以下依次设置左、右、上、下页面边距
    distance1 = Inches(0.2)
    distance2 = Inches(0)
    sec.left_margin = distance1
    sec.right_margin = distance1
    sec.top_margin = distance2
    sec.bottom_margin = distance2

    # 设置页面的宽度和高度(A4)
    sec.page_width = Inches(11.69)
    sec.page_height = Inches(8.27)

    flag_index = pd.Series(['', '', '', ''])  # 上一条运行的图片分类，用于下一次检测是否标题，换行等
    config_qty_tem = 1  # 用于检测是否大于3个config了，换页添加标题
    loop_count = 0

    #   循环逻辑：（1）Type不一样，Type标题+Config标题+图片，新建Par-RUN；（2）Config不一样，Config标题+图片，新建Rar-RUN；（3）正常绘图
    for i in pic_name.columns:

        # 主标题设置
        if pic_name[i][3] != flag_index[3]:
            if loop_count > 0:  # 第一页不换页
                document.add_page_break()
            page_title = document.add_heading("", level=1)
            page_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run_pt = page_title.add_run('\n\n\n' + 'TestData Plot Summary Report' + '\n' + pic_name[i][3])
            run_pt.font.name = u'Helvetica'
            run_pt.font.size = Pt(40.5)
            run_pt._element.rPr.rFonts.set(qn('w:eastAsia'), u'Helvetica')
            run_pt.font.color.rgb = RGBColor(0, 0, 0)
            run_pt.font.bold = True

        if pic_name[i][0] != flag_index[0] or (pic_name[i][1] != flag_index[1] and config_qty_tem % 3 == 0):  # 图类别
            document.add_page_break()
            page_title = document.add_heading('', level=1)
            page_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run_pt = page_title.add_run(pic_name[i][0])
            run_pt.font.name = u'Helvetica'
            run_pt.font.size = Pt(17)
            run_pt._element.rPr.rFonts.set(qn('w:eastAsia'), u'Helvetica')
            run_pt.font.color.rgb = RGBColor(0, 0, 0)
            run_pt.italic = True
            run_pt.font.bold = True

            page_title_2nd = document.add_heading('', level=2)
            page_title_2nd.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run_ct = page_title_2nd.add_run(pic_name[i][1])
            run_ct.font.name = u'Helvetica'
            run_ct.font.size = Pt(12)
            run_ct._element.rPr.rFonts.set(qn('w:eastAsia'), u'Helvetica')
            run_ct.font.color.rgb = RGBColor(0, 0, 0)

            pr_pic = document.add_paragraph()
            run_pict = pr_pic.add_run()

            run_pict.add_picture(pic_path + '\\' + i, height=Inches(2))
            flag_index = pic_name[i]

            config_qty_tem = 1


        elif pic_name[i][1] != flag_index[1]:  # Config

            page_title_2nd = document.add_heading('', level=2)
            page_title_2nd.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run_ct = page_title_2nd.add_run(pic_name[i][1])
            run_ct.font.name = u'Helvetica'
            run_ct.font.size = Pt(12)
            run_ct._element.rPr.rFonts.set(qn('w:eastAsia'), u'Helvetica')
            run_ct.font.color.rgb = RGBColor(0, 0, 0)

            pr_pic = document.add_paragraph()
            run_pict = pr_pic.add_run()

            run_pict.add_picture(pic_path + '\\' + i, height=Inches(2))
            flag_index = pic_name[i]

            config_qty_tem = config_qty_tem + 1

        else:
            run_pict.add_picture(pic_path + '\\' + i, height=Inches(2))
            flag_index = pic_name[i]

        loop_count = loop_count + 1

    config_name = entry_config.get()
    report_path = os.path.split(pic_path)[0] + '/' + config_name + f'_Data Summary Report_WD_{now_date}.docx'
    document.save(report_path)
    text_info_line = text_info.index(END).split('.')[0]
    text_info.insert(END, 'Finished Report in ' + report_path + '\n')
    text_info.tag_add('tag2', str(int(text_info_line) - 1) + '.0', text_info_line + '.0')  # Finish 打印突出显示
    text_info.tag_config('tag2', background='Cyan1', font=('Times', 10))
    text_info.see(END)
    text_info.update()


def run():
    time_t0 = time.time()
    text_info_line = 2  # 监控text_info已经输入到第几行
    text_info.delete(1.0, END)
    text_info.insert(1.0, 'Start -> Reading CSV File' + '\n')
    text_info.update()
    data_csv, USL, LSL, Overlay, Station, project = csv_process.csv_process(text_CSV.get())

    config_name = entry_config.get()
    data_ini = pd.read_excel(filepath_ini, sheet_name=Station)

    graph_flag = ''  # 画什么图
    graph_class = ''  # 图的分类标签

    boxplot_data_tem = pd.DataFrame()
    sweep_data_tem = pd.DataFrame()
    USL_tem = []
    LSL_tem = []
    boxplot_X_name = []

    print_info = ''  # 显示打印信息

    # 创建保存图片的文件夹
    global pic_dic
    filepath_csv = text_CSV.get()
    pic_dic = os.path.split(filepath_csv)[0] + '/' + 'Saved Photo' + '/'
    pic_dic_0 = os.path.split(filepath_csv)[0] + '/' + 'Saved Photo'
    isExists = os.path.exists(pic_dic)
    if not isExists:
        os.makedirs(pic_dic)
    v.set(pic_dic_0)

    out_data = pd.DataFrame(index=['Item', 'USL', 'LSL', 'Mean', 'STDEV', 'CPK'])  # 输出数据表格

    ini_pic_order = 0
    for i in range(len(data_ini['Data Item'])):

        # 画图种类定义+标签
        if data_ini.loc[i, 'Plot Type'] == data_ini.loc[i, 'Plot Type']:
            graph_flag = data_ini.loc[i, 'Plot Type']
            graph_class = data_ini.loc[i, 'Plot Title']
            ini_pic_order = ini_pic_order + 1

        item_name = data_ini.loc[i, 'Data Item']

        find_flag = 0
        for item_ext in LSL.index:
            if item_ext.find(item_name) >= 0:
                find_flag = 1
                break

        if find_flag:

            if graph_flag == 'Histogram' and data_ini.loc[i, 'Monitor'] == data_ini.loc[i, 'Monitor']:  # Monitor列判定是否画图

                # USL 配置文件有，则用配置文件，无则用数据USL
                if data_ini.loc[i, 'Upper Limit'] == data_ini.loc[i, 'Upper Limit']:
                    USL_histogram = data_ini.loc[i, 'Upper Limit']
                else:
                    USL_histogram = USL[item_name]

                # LSL 配置文件有，则用配置文件，无则用数据LSL
                if data_ini.loc[i, 'Lower Limit'] == data_ini.loc[i, 'Lower Limit']:
                    LSL_histogram = data_ini.loc[i, 'Lower Limit']
                else:
                    LSL_histogram = LSL[item_name]

                print_info, out_his = Histogram.histogram(Station, pic_dic, config_name, graph_class, item_name,
                                                          USL_histogram, LSL_histogram, data_csv[item_name],
                                                          data_ini.loc[i, 'Axis Upper Limit'],
                                                          data_ini.loc[i, 'Axis Lower Limit'], ini_pic_order,
                                                          entry_qty.get())
                out_data = pd.concat([out_data, out_his], axis=1)
                text_info.insert(END, print_info + '\n')
                text_info.see(END)
                text_info_line = text_info_line + 1
                text_info.update()

            elif graph_flag == 'Boxplot' and data_ini.loc[i, 'Monitor'] == data_ini.loc[i, 'Monitor']:  # Monitor列判定是否画图

                boxplot_data_tem = pd.concat([boxplot_data_tem, data_csv[item_name]], axis=1)

                if data_ini.loc[i, 'X-Axis Mark'] == data_ini.loc[i, 'X-Axis Mark']:
                    boxplot_X_name.append(data_ini.loc[i, 'X-Axis Mark'])

                # USL 配置文件有，则用配置文件，无则用数据USL
                if data_ini.loc[i, 'Upper Limit'] == data_ini.loc[i, 'Upper Limit']:
                    USL_tem.append(data_ini.loc[i, 'Upper Limit'])
                else:
                    USL_tem.append(USL[item_name])

                # LSL 配置文件有，则用配置文件，无则用数据LSL
                if data_ini.loc[i, 'Lower Limit'] == data_ini.loc[i, 'Lower Limit']:
                    LSL_tem.append(data_ini.loc[i, 'Lower Limit'])
                else:
                    LSL_tem.append(LSL[item_name])

                if data_ini.loc[i + 1, 'Plot Type'] == data_ini.loc[i + 1, 'Plot Type']:
                    print_info, out_box = Boxplot.Boxplot(Station, pic_dic, config_name, graph_class, USL_tem, LSL_tem,
                                                          boxplot_data_tem, data_ini.loc[i, 'Axis Upper Limit'],
                                                          data_ini.loc[i, 'Axis Lower Limit'], boxplot_X_name,
                                                          ini_pic_order, entry_qty.get())
                    boxplot_X_name = []
                    out_data = pd.concat([out_data, out_box], axis=1)
                    text_info.insert(END, print_info + '\n')
                    text_info_line = text_info_line + 1
                    text_info.see(END)
                    text_info.update()
                    boxplot_data_tem = pd.DataFrame()
                    USL_tem = []
                    LSL_tem = []

            elif graph_flag == 'Sweep' and data_ini.loc[i, 'Monitor'] == data_ini.loc[i, 'Monitor']:  # Monitor列判定是否画图
                for j in data_csv.columns:
                    if j.find(item_name) >= 0:
                        sweep_data_tem = pd.concat([sweep_data_tem, data_csv[j]], axis=1)
                        USL_tem.append(USL[j])
                        LSL_tem.append(LSL[j])
                print_info = Sweep.Sweep(Station, pic_dic, config_name, graph_class, USL_tem, LSL_tem, sweep_data_tem,
                                         data_ini.loc[i, 'Axis Upper Limit'], data_ini.loc[i, 'Axis Lower Limit'],
                                         ini_pic_order, entry_qty.get())
                text_info.insert(END, print_info + '\n')
                text_info_line = text_info_line + 1
                text_info.see(END)
                text_info.update()
                sweep_data_tem = pd.DataFrame()
                USL_tem = []
                LSL_tem = []

        else:
            print_info = f'Error -> {item_name} IS not Found on Data !'
            text_info.insert(END, print_info + '\n')
            text_info.tag_add('tag0', str(text_info_line) + '.9',
                              str(text_info_line) + '.' + str(9 + len(item_name)))  # Error 打印突出显示
            text_info.tag_config('tag0', background='red', font=('Times'))
            text_info_line = text_info_line + 1
            text_info.see(END)
            text_info.update()

    # 保存数据到Excel Table
    save_table_name = project + '_' + Station + f'_Summary Table_{now_date}.xlsx'
    out_data_path = os.path.split(filepath_csv)[0] + '/' + save_table_name
    app = xw.App(visible=False, add_book=False)
    if os.path.exists(out_data_path):
        wb = app.books.open(out_data_path)
        for i in wb.sheets:
            if i.name == entry_config.get():
                wb.sheets[entry_config.get()].delete()
        ws = wb.sheets.add(entry_config.get())
        ws.range('A1').expand('table').value = out_data
        ws.range('A1').api.EntireRow.Delete()
    else:
        wb = app.books.add()
        wb.sheets[0].name = entry_config.get()
        ws = wb.sheets[0]
        ws.range('A1').expand('table').value = out_data
        ws.range('A1').api.EntireRow.Delete()
    wb.save(out_data_path)
    wb.close()
    app.quit()

    time_delta = time.time() - time_t0
    text_info.insert(END, 'Finished All in ' + str(round(time_delta, 1)) + ' Seconds' + '\n')
    text_info.tag_add('tag', str(text_info_line) + '.0', str(text_info_line + 1) + '.0')  # Finish 打印突出显示
    text_info.tag_config('tag', background='green', font=('Times', 15))
    text_info.see(END)
    text_info.update()
    text_info_line = text_info_line + 1


# 主窗口
root = Tk()
root.title('AAC DataPlot Summary_1.5')
root.resizable(0, 0)
root.geometry('700x500')

y_start = 0.01
y_interval = 0.07

# Config label
entry_config = Entry(root, font=('Times', 15, 'bold'))
entry_config.place(relx=0.01, rely=y_start, relwidth=0.36, relheight=0.06)
lbe_config = Label(root, text='CONFIG')
lbe_config.place(relx=0.38, rely=y_start, relwidth=0.11, relheight=0.06)

# Qty label
entry_qty = Entry(root, font=('Times', 15, 'bold'))
entry_qty.place(relx=0.51, rely=y_start, relwidth=0.36, relheight=0.06)
lbe_qty = Label(root, text='N Qty')
lbe_qty.place(relx=0.88, rely=y_start, relwidth=0.11, relheight=0.06)

# CSV 文件输入
text_CSV = Entry(root, font=('Helvetica', 10), relief=GROOVE)
text_CSV.place(relx=0.01, rely=y_start + y_interval, relwidth=0.86, relheight=0.06)
btn_CSV = Button(root, text='CSV Path', command=run_csv)
btn_CSV.place(relx=0.88, rely=y_start + y_interval, relwidth=0.11, relheight=0.06)

# 配置文件输入
text_ini = Entry(root, font=('Helvetica', 10))
text_ini.place(relx=0.01, rely=y_start + 2 * y_interval, relwidth=0.86, relheight=0.06)
btn_ini = Button(root, text='INI File', command=run_ini)
btn_ini.place(relx=0.88, rely=y_start + 2 * y_interval, relwidth=0.11, relheight=0.06)

# 图片路径
v = StringVar()
text_pic = Label(root, justify="left", font=('Helvetica', 10), relief=GROOVE, textvariable=v)
text_pic.place(relx=0.01, rely=y_start + 3 * y_interval, relwidth=0.86, relheight=0.06)
lbe_pic = Button(root, text='PIC Path', command=run_pic)
lbe_pic.place(relx=0.88, rely=y_start + 3 * y_interval, relwidth=0.11, relheight=0.06)

# 信息输出框 & 执行按钮
text_info = Text(root, font=('Times', 10))
text_info.place(relx=0.01, rely=y_start + 4 * y_interval, relwidth=0.86, relheight=0.695)

# 信息输出框的滚动条
scroll = Scrollbar()
# 放到窗口的右侧, 填充Y竖直方向
scroll.place(relx=0.8492, rely=y_start + 4 * y_interval + 0.001, relwidth=0.02, relheight=0.695 - 0.002)
scroll.config(command=text_info.yview)
text_info.config(yscrollcommand=scroll.set)
btn_run = Button(root, text='START', command=run)
btn_run.place(relx=0.88, rely=y_start + 4 * y_interval, relwidth=0.11, relheight=0.06)
btn_report = Button(root, text='REPORT', command=run_report)
btn_report.place(relx=0.88, rely=y_start + 5 * y_interval, relwidth=0.11, relheight=0.06)

dt = datetime.now()
now_date = dt.strftime('%Y%m%d')
overdue_date = '20230704'

if now_date > overdue_date:
    tkinter.messagebox.askokcancel(title='Error', message='License Expired !')
    exit()

root.mainloop()  # 进入消息循环
